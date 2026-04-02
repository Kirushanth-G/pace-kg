"""
Pipeline runner that wraps all 9 steps from edu_kg_gliner.py.
Each step function is extracted from the reference Colab implementation.
"""

import json
import os
import re
import time
from dataclasses import asdict, dataclass, field, replace
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import fitz
import pytesseract
import spacy
from docx import Document
from docx.shared import Pt
from gliner import GLiNER
from langchain.schema import HumanMessage, SystemMessage
from langchain_groq import ChatGroq
from marker.converters.pdf import PdfConverter
from marker.models import create_model_dict
from neo4j import GraphDatabase
from PIL import Image
from sentence_transformers import SentenceTransformer, util

from core.config import settings


# ══════════════════════════════════════════════════════════════════════════════
# Data Models
# ══════════════════════════════════════════════════════════════════════════════


@dataclass
class SlideMarkdown:
    slide_id: str
    page_number: int
    raw_markdown: str
    doc_id: str


@dataclass
class SlideContent:
    slide_id: str
    page_number: int
    doc_id: str
    headings: List[str] = field(default_factory=list)
    body_text: List[str] = field(default_factory=list)
    bullets: List[str] = field(default_factory=list)
    table_cells: List[str] = field(default_factory=list)
    captions: List[str] = field(default_factory=list)
    code_lines: List[str] = field(default_factory=list)
    heading_phrases: List[str] = field(default_factory=list)
    clean_text: str = ""
    ocr_applied: bool = False


@dataclass
class Keyphrase:
    phrase: str
    score: float
    source_type: str  # heading | bullet | table | caption | body | injected
    slide_id: str
    doc_id: str
    appears_in: str


@dataclass
class Triple:
    subject: str
    relation: str
    object: str
    evidence: str
    confidence: float
    slide_id: str
    doc_id: str
    source: str = "extraction"


@dataclass
class ConceptNode:
    name: str
    aliases: List[str]
    slide_ids: List[str]
    source_type: str
    keyphrase_score: float
    final_weight: float
    doc_id: str
    needs_review: bool = False


@dataclass
class ExpansionEdge:
    subject: str
    relation: str  # always "relatedConcept"
    object: str
    source: str  # always "expansion"
    confidence: float
    slide_id: str
    doc_id: str


@dataclass
class SlideSummary:
    slide_id: str
    page_number: int
    heading: str
    summary: str
    key_terms: List[str]
    doc_id: str


# ══════════════════════════════════════════════════════════════════════════════
# Pipeline Runner Class
# ══════════════════════════════════════════════════════════════════════════════


class PipelineRunner:
    """Runs the complete 9-step PACE-KG pipeline on a PDF."""

    def __init__(self, doc_id: str, pdf_path: str, output_dir: str):
        self.doc_id = doc_id
        self.pdf_path = pdf_path
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Models (loaded once, reused across steps)
        self.marker_models = None
        self.gliner = None
        self.nlp = None
        self.minilm = None
        self.sbert = None
        self.llm_primary = None
        self.llm_fallback = None
        self.llm_summary = None

        # Set Groq API key in environment
        os.environ["GROQ_API_KEY"] = settings.groq_api_key

    def _save_json(self, data, step_name: str):
        """Save step output as JSON."""
        path = self.output_dir / f"{self.doc_id}_{step_name}.json"
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return path

    def _load_json(self, step_name: str):
        """Load step output from JSON."""
        path = self.output_dir / f"{self.doc_id}_{step_name}.json"
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 1 — Marker PDF Parsing
    # ══════════════════════════════════════════════════════════════════════════

    def step1_parse_pdf(self) -> List[SlideMarkdown]:
        """Parse PDF using Marker + OCR fallback."""
        print(f"[Step 1] Parsing PDF: {self.pdf_path}")

        # Load Marker models once
        if self.marker_models is None:
            print("Loading Marker models (first run ~5min download)...")
            self.marker_models = create_model_dict()

        # Run Marker with paginate_output enabled via config dict
        converter = PdfConverter(
            artifact_dict=self.marker_models,
            config={"paginate_output": True, "page_separator": settings.page_sep},
        )
        rendered = converter(self.pdf_path)
        full_md = rendered.markdown

        # Split into pages using the separator
        page_mds = [p.strip() for p in full_md.split(settings.page_sep)]
        slides_md: List[SlideMarkdown] = []

        # Track seen content for duplicate detection
        seen_content: List[set] = []

        for i, md in enumerate(page_mds, start=1):
            # Check if empty
            if self._is_empty_page(md):
                # OCR fallback
                md_ocr = self._ocr_page(i)
                if md_ocr:
                    md = md_ocr
                else:
                    continue  # Skip truly empty pages

            # Duplicate detection
            words = set(re.findall(r"\w+", md.lower()))
            if self._is_duplicate(words, seen_content):
                print(f"  Slide {i}: duplicate detected, skipping")
                continue

            seen_content.append(words)
            slide = SlideMarkdown(
                slide_id=f"slide_{i:03d}",
                page_number=i,
                raw_markdown=md,
                doc_id=self.doc_id,
            )
            slides_md.append(slide)

        # Save output
        data = [asdict(s) for s in slides_md]
        self._save_json(data, "step1_parsed")
        print(f"[Step 1] Parsed {len(slides_md)} slides")
        return slides_md

    def _is_empty_page(self, markdown: str) -> bool:
        """Check if page is empty (FIX S1-2: strips {N} tags)."""
        cleaned = markdown.strip()
        cleaned = re.sub(r"\{\d+\}", "", cleaned).strip()
        return len(cleaned) < 10

    def _ocr_page(self, page_num: int) -> str:
        """OCR fallback for empty pages."""
        try:
            doc = fitz.open(self.pdf_path)
            page = doc[page_num - 1]
            zoom = 3
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)
            doc.close()

            if text.strip():
                lines = [
                    f"> OCR: {line}"
                    for line in text.strip().split("\n")
                    if line.strip()
                ]
                return "\n".join(lines)
        except Exception as e:
            print(f"  OCR failed on page {page_num}: {e}")
        return ""

    def _is_duplicate(
        self, words: set, seen: List[set], threshold=0.5, window=2
    ) -> bool:
        """Jaccard similarity duplicate detection (FIX S1-1)."""
        if len(words) < 5:
            return False
        for prev in seen[-window:]:
            if len(prev) < 5:
                continue
            jaccard = len(words & prev) / len(words | prev)
            if jaccard >= threshold:
                return True
        return False

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 2 — Markdown Preprocessor
    # ══════════════════════════════════════════════════════════════════════════

    def step2_preprocess(self, slides_md: List[SlideMarkdown]) -> List[SlideContent]:
        """Preprocess markdown into structured SlideContent."""
        print(f"[Step 2] Preprocessing {len(slides_md)} slides")

        slides_content: List[SlideContent] = []
        for sm in slides_md:
            sc = self._parse_slide(sm)
            slides_content.append(sc)

        # Cross-slide repetition filter
        if len(slides_content) >= 5:
            self._remove_common_blocks(slides_content)

        # Save output
        data = [asdict(s) for s in slides_content]
        self._save_json(data, "step2_preprocessed")
        print(f"[Step 2] Preprocessed {len(slides_content)} slides")
        return slides_content

    def _parse_slide(self, sm: SlideMarkdown) -> SlideContent:
        """Parse slide markdown into typed buckets."""
        sc = SlideContent(
            slide_id=sm.slide_id,
            page_number=sm.page_number,
            doc_id=sm.doc_id,
            ocr_applied="> OCR:" in sm.raw_markdown,
        )

        lines = sm.raw_markdown.split("\n")
        in_code_fence = False

        for line in lines:
            stripped = line.strip()

            # Code fence toggle
            if stripped.startswith("```"):
                in_code_fence = not in_code_fence
                continue

            # Noise filter
            if self._is_noise(stripped):
                continue

            if in_code_fence:
                # Inside code fence: check for prose bullets
                if re.match(r"^[-*•◦▪▫–—►◮]\s", stripped):
                    sc.bullets.append(re.sub(r"^[-*•◦▪▫–—►◮]\s+", "", stripped))
                elif self._is_code_line(stripped):
                    sc.code_lines.append(stripped)
                continue

            # Structural routing
            if stripped.startswith("#"):
                heading = re.sub(r"^#+\s*", "", stripped)
                sc.headings.append(heading)
                sc.heading_phrases.append(heading)
            elif stripped.startswith("> OCR:"):
                content = stripped[6:].strip()
                self._classify_ocr_line(content, sc)
            elif stripped.startswith(">"):
                sc.captions.append(stripped[1:].strip())
            elif re.match(r"^[-*•◦▪▫–—►◮]\s", stripped):
                sc.bullets.append(re.sub(r"^[-*•◦▪▫–—►◮]\s+", "", stripped))
            elif "|" in stripped and stripped.count("|") >= 2:
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                sc.table_cells.extend(cells)
            elif stripped:
                sc.body_text.append(stripped)

        # Assemble clean_text
        sc.clean_text = " ".join(
            sc.headings + sc.body_text + sc.bullets + sc.table_cells + sc.captions
        )
        return sc

    def _is_noise(self, text: str) -> bool:
        """Check if line matches noise patterns."""
        patterns = [
            r"^\s*page\s+\d+\s*(of\s+\d+)?\s*$",
            r"^\s*\d+\s*$",
            r"^(?:©|\(c\)|copyright\b)",
            r"^\s*(references|bibliography)\s*$",
            r"https?://\S+",
            r"^\s*\[\d+\]",
            r"^!\[.*?\]\(.*?\)$",
            r"^quiz\s*:",
            r"^\{\d+\}$",
            r"^(input|output|legend|infrastructure)(\s+replaced.*|\s+kept.*|\s+storage.*)?$",
        ]
        return any(re.match(p, text, re.IGNORECASE) for p in patterns)

    def _is_code_line(self, text: str) -> bool:
        """Heuristic code line detector."""
        code_chars = set("{}();=[]")
        n_code = sum(1 for c in text if c in code_chars)
        words = text.split()

        # Rule 1: function call
        if "(" in text and ")" in text and n_code >= 3 and len(words) <= 8:
            return True

        # Rule 2: camelCase with few commas
        if re.search(r"[a-z][A-Z]", text) and text.count(",") < 2:
            return True

        # Rule 3: dot-access (new FIX S2-6)
        if re.match(r"^\w+\.\w+(\.\w+)*$", text):
            return True

        return False

    def _classify_ocr_line(self, content: str, sc: SlideContent):
        """Re-classify OCR lines."""
        if content.startswith("#"):
            sc.headings.append(content[1:].strip())
        elif re.match(r"^[-*•]\s", content):
            sc.bullets.append(re.sub(r"^[-*•]\s+", "", content))
        elif self._is_code_line(content):
            sc.code_lines.append(content)
        else:
            sc.body_text.append(content)

    def _remove_common_blocks(self, slides: List[SlideContent]):
        """Remove blocks appearing on >50% of slides."""
        block_counts: Dict[str, int] = {}
        for sc in slides:
            blocks = set(sc.headings + sc.body_text + sc.bullets)
            for b in blocks:
                block_counts[b] = block_counts.get(b, 0) + 1

        threshold = len(slides) * 0.5
        common = {b for b, cnt in block_counts.items() if cnt > threshold}

        if common:
            print(f"  Removing {len(common)} common blocks")
            for sc in slides:
                sc.headings = [h for h in sc.headings if h not in common]
                sc.body_text = [b for b in sc.body_text if b not in common]
                sc.bullets = [b for b in sc.bullets if b not in common]
                sc.clean_text = " ".join(
                    sc.headings
                    + sc.body_text
                    + sc.bullets
                    + sc.table_cells
                    + sc.captions
                )

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 3 — Keyphrase Extraction (GLiNER + LLM fallback)
    # ══════════════════════════════════════════════════════════════════════════

    def step3_extract_keyphrases(
        self, slides: List[SlideContent]
    ) -> Dict[str, List[Keyphrase]]:
        """Extract keyphrases using GLiNER + LLM fallback."""
        print(f"[Step 3] Extracting keyphrases from {len(slides)} slides")

        # Load models once
        if self.gliner is None:
            print("Loading GLiNER large-v2.1...")
            self.gliner = GLiNER.from_pretrained("urchade/gliner_large-v2.1")
        if self.nlp is None:
            print("Loading spaCy...")
            self.nlp = spacy.load("en_core_web_sm")
        if self.minilm is None:
            print("Loading MiniLM...")
            self.minilm = SentenceTransformer("all-MiniLM-L6-v2")
        if self.llm_fallback is None:
            # Use 70B model for Step 3 fallback (higher quality keyphrases)
            self.llm_fallback = ChatGroq(
                model=settings.llm_primary,  # llama-3.3-70b-versatile
                temperature=0,
                api_key=settings.groq_api_key,
                model_kwargs={"response_format": {"type": "json_object"}},
            )

        labels = [
            "Academic Concept",
            "Theoretical Principle",
            "Technical Term",
            "Process or Method",
            "System or Framework",
            "Formula or Equation",
        ]

        all_keyphrases: Dict[str, List[Keyphrase]] = {}

        for sc in slides:
            # Pedagogical filter
            word_count = len(sc.clean_text.split())
            if sc.page_number == 1 and word_count < 40:
                continue
            if word_count < 8:
                continue

            # Build extract_text (headings first, no code)
            extract_text = " ".join(
                sc.headings + sc.body_text + sc.bullets + sc.table_cells + sc.captions
            )
            if not extract_text.strip():
                continue

            # GLiNER extraction
            entities = self.gliner.predict_entities(
                extract_text, labels, threshold=settings.gliner_threshold
            )
            kps = []

            for ent in entities:
                phrase = ent["text"].lower().strip()
                score = ent["score"]
                source_type = self._assign_source_type(phrase, sc)
                appears_in = self._find_sentence(phrase, sc.clean_text)

                kp = Keyphrase(
                    phrase=phrase,
                    score=score,
                    source_type=source_type,
                    slide_id=sc.slide_id,
                    doc_id=sc.doc_id,
                    appears_in=appears_in,
                )
                kps.append(kp)

            # LLM fallback if GLiNER returns ≤1 phrase
            if len(kps) <= 1 and word_count >= 20:
                llm_kps = self._llm_keyphrase_fallback(extract_text, sc)
                kps.extend(llm_kps)

            # Post-processing
            kps = self._collapse_duplicates(kps)
            kps = [k for k in kps if len(k.phrase) >= 3]

            # Heading boost
            for kp in kps:
                if kp.source_type == "heading":
                    kp.score = min(kp.score + settings.heading_boost, 1.0)

            # Deduplicate near-synonyms
            kps = self._deduplicate_keyphrases(kps)

            # Sort and cap
            kps.sort(key=lambda k: k.score, reverse=True)
            kps = kps[: settings.keyphrase_max]

            all_keyphrases[sc.slide_id] = kps

        # Save output
        data = {sid: [asdict(k) for k in kps] for sid, kps in all_keyphrases.items()}
        self._save_json(data, "step3_keyphrases")
        print(f"[Step 3] Extracted keyphrases for {len(all_keyphrases)} slides")
        return all_keyphrases

    def _assign_source_type(self, phrase: str, sc: SlideContent) -> str:
        """Determine source type of keyphrase."""
        pl = phrase.lower()
        for h in sc.headings:
            if pl in h.lower():
                return "heading"
        for b in sc.bullets:
            if pl in b.lower():
                return "bullet"
        for t in sc.table_cells:
            if pl in t.lower():
                return "table"
        for c in sc.captions:
            if pl in c.lower():
                return "caption"
        return "body"

    def _find_sentence(self, phrase: str, text: str) -> str:
        """Find sentence containing phrase."""
        if not text:
            return ""
        doc = self.nlp(text)
        pl = phrase.lower()
        for sent in doc.sents:
            if pl in sent.text.lower():
                return sent.text.strip()
        return text[:200]

    def _llm_keyphrase_fallback(self, text: str, sc: SlideContent) -> List[Keyphrase]:
        """LLM fallback for sparse slides - uses reference implementation prompt."""
        system_prompt = """You are extracting educational keyphrases from a single lecture slide.

Return ONLY a JSON object: {"keyphrases": ["phrase1", "phrase2", ...]}

Rules:
- Maximum 12 phrases.
- Only include named concepts, technical terms, algorithms, data structures,
  or named methods that a student MUST understand from this slide.
- Good examples: "bubble sort", "ArrayList", "dynamic resizing", "FIFO",
  "binarySearch", "hashmap", "null reference", "linked list"
- Do NOT include: generic words (value, type, size, element, method),
  instructional phrases (to construct, any list of, storing a value),
  professor names, university names, course codes, dates, or slide titles.
- If the slide has no educational content, return {"keyphrases": []}
- Return ONLY the JSON object. No explanation."""

        user_prompt = f"""SLIDE TEXT:
{text}"""

        try:
            from langchain.schema import HumanMessage, SystemMessage

            resp = self.llm_fallback.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt),
                ]
            )
            raw = json.loads(resp.content)
            phrases = raw.get("keyphrases", [])
            if not isinstance(phrases, list):
                return []

            kps = []
            for phrase in phrases:
                phrase = phrase.lower().strip()
                if len(phrase) >= 3:
                    kp = Keyphrase(
                        phrase=phrase,
                        score=settings.llm_fallback_score,
                        source_type=self._assign_source_type(phrase, sc),
                        slide_id=sc.slide_id,
                        doc_id=sc.doc_id,
                        appears_in=self._find_sentence(phrase, sc.clean_text),
                    )
                    kps.append(kp)
            return kps
        except Exception as e:
            print(f"  LLM fallback failed: {e}")
            return []

    def _collapse_duplicates(self, kps: List[Keyphrase]) -> List[Keyphrase]:
        """Keep highest score per unique phrase."""
        seen: Dict[str, Keyphrase] = {}
        for kp in kps:
            if kp.phrase not in seen or kp.score > seen[kp.phrase].score:
                seen[kp.phrase] = kp
        return list(seen.values())

    def _deduplicate_keyphrases(self, kps: List[Keyphrase]) -> List[Keyphrase]:
        """Remove near-duplicate keyphrases using SBERT."""
        if len(kps) <= 1:
            return kps

        sorted_kps = sorted(kps, key=lambda k: k.score, reverse=True)
        kept: List[Keyphrase] = []

        for kp in sorted_kps:
            if not kept:
                kept.append(kp)
                continue

            # Check similarity with kept phrases
            emb_new = self.minilm.encode(kp.phrase, convert_to_tensor=True)
            is_dup = False

            for existing in kept:
                emb_ex = self.minilm.encode(existing.phrase, convert_to_tensor=True)
                sim = util.cos_sim(emb_new, emb_ex).item()
                if sim >= settings.dedup_sim_threshold:
                    is_dup = True
                    break

            if not is_dup:
                kept.append(kp)

        return kept

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 4 — LLM Triple Extraction (continues in next chunk...)
    # ══════════════════════════════════════════════════════════════════════════

    def step4_extract_triples(
        self, slides: List[SlideContent], keyphrases: Dict[str, List[Keyphrase]]
    ) -> List[Triple]:
        """Step 4: LLM triple extraction with 4-layer validation."""
        print(f"[Step 4] Extracting triples from {len(slides)} slides")

        # Load SBERT once
        if self.sbert is None:
            print("Loading SBERT all-mpnet-base-v2...")
            self.sbert = SentenceTransformer("all-mpnet-base-v2")

        all_triples: List[Triple] = []
        prev_keyphrase: Optional[Keyphrase] = None
        rate_limit_skipped = 0

        for idx, sc in enumerate(slides, 1):
            kps = keyphrases.get(sc.slide_id, [])
            if len(kps) < 2:
                continue

            # Cross-slide enrichment: inject top keyphrase from prev slide
            original_kps = [k for k in kps if k.source_type != "injected"]
            if prev_keyphrase and prev_keyphrase.phrase not in [k.phrase for k in kps]:
                injected = replace(
                    prev_keyphrase,
                    slide_id=sc.slide_id,
                    score=0.35,
                    source_type="injected",
                )
                kps.append(injected)

            # Check if LLM slide (all original kps have score 0.40)
            is_llm_slide = len(original_kps) > 0 and all(
                k.score == settings.llm_fallback_score for k in original_kps
            )

            # Extract triples from LLM
            print(f"  Processing slide {idx}/{len(slides)}: {sc.slide_id}")
            triples = self._extract_triples_for_slide(sc, kps, is_llm_slide)
            if triples is None:  # Rate limit indicator
                rate_limit_skipped += 1
                continue
            all_triples.extend(triples)

            # Update prev_keyphrase
            if original_kps:
                prev_keyphrase = max(original_kps, key=lambda k: k.score)

        # Save output
        data = [asdict(t) for t in all_triples]
        self._save_json(data, "step4_triples")

        if rate_limit_skipped > 0:
            print(
                f"[Step 4] WARNING: Skipped {rate_limit_skipped} slides due to rate limits"
            )
            print(
                f"[Step 4] Extracted {len(all_triples)} triples from {len(slides) - rate_limit_skipped} slides"
            )
            print(
                f"[Step 4] RECOMMENDATION: Wait ~8 hours for limit reset OR upgrade to Dev Tier ($0.10/million tokens)"
            )
        else:
            print(f"[Step 4] Extracted {len(all_triples)} triples")

        return all_triples

    def _extract_triples_for_slide(
        self, sc: SlideContent, kps: List[Keyphrase], is_llm_slide: bool
    ) -> Optional[List[Triple]]:
        """Extract and validate triples for a single slide.

        Returns:
            List[Triple]: Extracted triples (may be empty if no valid triples found)
            None: If rate-limited (signals to caller to skip this slide)
        """

        # Build prompt
        system_prompt = self._get_system_prompt()
        kp_list = ", ".join([k.phrase for k in kps])
        heading = sc.headings[0] if sc.headings else sc.slide_id

        user_prompt = f"""SLIDE TITLE: {heading}

KEYPHRASE LIST:
{kp_list}

SLIDE TEXT:
{sc.clean_text}

Return triples as JSON array with this format:
[
  {{
    "subject": "concept1",
    "relation": "isDefinedAs",
    "object": "concept2",
    "evidence": "exact sentence from slide text",
    "confidence": 0.85
  }}
]"""

        # Call LLM (with fallback on 429)
        triples_raw = []
        try:
            resp = self.llm_primary.invoke(
                [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=user_prompt),
                ]
            )
            triples_raw = json.loads(resp.content)
        except Exception as e:
            error_str = str(e)
            # Check for rate limit error (HTTP 429)
            if "429" in error_str or "rate_limit" in error_str.lower():
                print(f"  Primary model (70B) rate limited, trying fallback (8B)...")
                time.sleep(2)
                try:
                    resp = self.llm_fallback.invoke(
                        [
                            SystemMessage(content=system_prompt),
                            HumanMessage(content=user_prompt),
                        ]
                    )
                    triples_raw = json.loads(resp.content)
                except Exception as fallback_error:
                    # If fallback also fails, check if it's also rate limited
                    fallback_error_str = str(fallback_error)
                    if (
                        "429" in fallback_error_str
                        or "rate_limit" in fallback_error_str.lower()
                    ):
                        print(
                            f"  ⚠️  Both models rate limited for {sc.slide_id}. Skipping slide."
                        )
                        return None  # Signal rate limit to caller
                    else:
                        print(f"  Fallback model failed: {fallback_error}")
                        return []
            else:
                print(f"  LLM extraction failed: {e}")
                return []

        # Parse and validate
        triples: List[Triple] = []
        for t in triples_raw:
            triple = Triple(
                subject=t["subject"].lower().strip(),
                relation=t["relation"],
                object=t["object"].lower().strip(),
                evidence=t["evidence"],
                confidence=t["confidence"],
                slide_id=sc.slide_id,
                doc_id=sc.doc_id,
                source="extraction",
            )

            if self._validate_triple(triple, kps, sc, is_llm_slide):
                triples.append(triple)

        # Dedup
        triples = self._dedup_triples(triples)
        return triples

    def _get_system_prompt(self) -> str:
        """Return the LLM system prompt for triple extraction."""
        return """You are a knowledge graph construction assistant for educational materials.

STRICT RULES -- violating any rule invalidates your entire response:
1. ONLY use concepts from the KEYPHRASE LIST as subject and object. Nothing else.
2. Every triple MUST include an EXACT sentence copied from the SLIDE TEXT as evidence.
3. If no supporting sentence exists in the slide text, omit that triple entirely.
4. Use ONLY these 8 relation types. Follow the DIRECTION rules exactly:

   isPrerequisiteOf   subject must be understood BEFORE object can be understood
                      e.g. "sort --[isPrerequisiteOf]--> binarysearch"
                      (you must sort an array before binary search works on it)

   isDefinedAs        subject IS the concept being defined; object is its definition
                      e.g. "queue --[isDefinedAs]--> fifo"
                      NEVER reverse: concept first, definition second
                      ONE definition per concept — do not repeat with synonyms

   isExampleOf        subject is a SPECIFIC INSTANCE or IMPLEMENTATION of object
                      e.g. "arraylist --[isExampleOf]--> list"
                      e.g. "hashmap --[isExampleOf]--> map"
                      e.g. "bubble sort --[isExampleOf]--> sorting algorithms"
                      USE THIS when X is a named implementation of an interface
                      or a named member of a broader category

   contrastedWith     subject AND object are EXPLICITLY compared or contrasted
                      e.g. "hashmap --[contrastedWith]--> treemap"
                      USE THIS when: slide title has "vs"/"comparison"/"advantages",
                      OR slide text uses "unlike", "whereas", "faster than",
                      "compared to", "better for"

   appliedIn          subject IS USED IN, IS PASSED TO, or OPERATES INSIDE object
                      e.g. "arrays --[appliedIn]--> method"
                      e.g. "comparable --[appliedIn]--> sorting"
                      USE THIS when slide says "X can be passed to Y",
                      "X is used in Y", "X works inside Y"

   isPartOf           subject is a PHYSICAL or STRUCTURAL FIELD/COMPONENT of object
                      e.g. "link --[isPartOf]--> node"
                      e.g. "length --[isPartOf]--> array object"
                      ONLY for literal structural parts: fields of a class,
                      components of a data structure, constants of an object.
                      DO NOT use for: passing/usage, implementations, categories

   causeOf            subject DIRECTLY CAUSES object as a result
                      e.g. "off-by-one error --[causeOf]--> arrayindexoutofboundsexception"

   isGeneralizationOf subject is the BROADER CATEGORY; object is a narrower member
                      e.g. "sorting algorithms --[isGeneralizationOf]--> bubble sort"
                      e.g. "map --[isGeneralizationOf]--> hashmap"
                      The BROADER concept is ALWAYS the subject.

5. CRITICAL DISAMBIGUATION — read every rule before choosing a relation:

   PASSING / USAGE (use appliedIn, NOT isPartOf):
   - "X can be passed to Y"  →  X appliedIn Y
   - "X is used in Y"        →  X appliedIn Y
   - "X works with Y"        →  X appliedIn Y
   - NEVER use isPartOf when the slide is describing passing or usage

   METHOD / OPERATION (use isPartOf for ownership, NOT isExampleOf):
   - If X is a METHOD or OPERATION name (add, remove, enqueue, search):
     → X isPartOf the class/structure it belongs to
     e.g. "add isPartOf arraylist",  "enqueue isPartOf queue"
   - NEVER use isExampleOf for a method or operation name

   IMPLEMENTATION (use isExampleOf, NOT isPartOf):
   - If X is a NAMED IMPLEMENTATION of an interface or abstract concept:
     → X isExampleOf the interface/concept
     e.g. "hashmap isExampleOf map",  "treeset isExampleOf set interface"

   ACRONYM / EXPANSION (only one definition triple per concept):
   - If the keyphrase list contains BOTH an acronym AND its expansion
     (e.g. "fifo" and "first in, first out"), they refer to the SAME concept.
   - Produce ONE isDefinedAs triple using the ACRONYM as subject:
     e.g. "queue isDefinedAs fifo"
   - Do NOT also produce "queue isDefinedAs first in, first out"
   - Do NOT produce isGeneralizationOf or isExampleOf between the acronym
     and its expansion — they are the same concept, not two different ones

   isDefinedAs direction:
   - The NAMED CONCEPT is always the subject.
   - The DEFINITION or DESCRIPTION is the object.
   - "array isDefinedAs ordered list" ✓
   - "ordered list isDefinedAs array" ✗

   isGeneralizationOf direction:
   - The BROADER CATEGORY is the subject, the NARROWER member is the object.
   - "sorting algorithms isGeneralizationOf bubble sort" ✓
   - "bubble sort isGeneralizationOf sorting algorithms" ✗

   Structural membership:
   - If slide says "X consists of Y" / "X has Y" / "X contains Y":
     → Y isPartOf X
   - If slide says "X is a type of Y" / "X is an implementation of Y":
     → X isExampleOf Y

6. DIVERSITY REQUIREMENT:
   - Consider ALL 8 relation types before choosing
   - Do NOT use isPartOf for more than 40% of triples in this response
   - If the slide title contains "vs", "comparison", or "advantages",
     produce at least one contrastedWith triple

7. SLIDE HEADING RULE:
   - The slide title is a topic label, not an educational concept
   - Do NOT use the full slide title as subject or object
   - Use the specific concepts taught within the slide

8. STRUCTURAL LABEL RULE:
   - Do NOT use instructional sub-headings as subject or object:
     e.g. "construction", "storing a value", "to construct an array list"
   - If the keyphrase list contains only such labels, return []

9. DUPLICATE RULE:
   - Do NOT produce two triples with the same (subject, object) pair
   - If you find two valid relations for the same pair, keep only the
     most specific / highest-confidence one
   - Do NOT produce both X isGeneralizationOf Y and Y isExampleOf X
     for the same pair — pick ONE direction

10. Return ONLY a valid JSON array. No markdown, no explanation, no preamble.
11. If no valid triples found, return: []

CRITICAL CONSTRAINTS:
1. NO ADMINISTRATIVE ENTITIES: No people names, universities, course codes, dates.
2. PEDAGOGICAL ONLY: Only academic/domain concepts taught in the material.
3. EMPTY FALLBACK: Title slides, admin slides, ToC slides → return []"""

    def _is_structural_label(self, phrase: str) -> bool:
        """Return True if phrase is an instructional slide label, not a concept."""
        p = phrase.strip().lower()
        _starters = (
            "to ",
            "any ",
            "how ",
            "what ",
            "when ",
            "where ",
            "why ",
            "see ",
            "let ",
            "note ",
            "given ",
            "recall ",
            "consider ",
        )
        if any(p.startswith(s) for s in _starters) and len(p.split()) >= 4:
            return True
        # Gerund-led instructional phrases (e.g. "storing a value", "declaring and using")
        words = p.split()
        if words and words[0].endswith("ing") and len(words) >= 3:
            return True
        return False

    def _validate_triple(
        self, triple: Triple, kps: List[Keyphrase], sc: SlideContent, is_llm_slide: bool
    ) -> bool:
        """3-layer validation (ALL must pass) - matching reference implementation."""

        # Layer 1: Anchor check
        kp_phrases = {k.phrase for k in kps}
        subj_clean = triple.subject.replace("(", "").replace(")", "").strip()
        obj_clean = triple.object.replace("(", "").replace(")", "").strip()

        if subj_clean not in kp_phrases or obj_clean not in kp_phrases:
            return False
        if triple.subject == triple.object:
            return False
        if triple.relation not in [
            "isPrerequisiteOf",
            "isDefinedAs",
            "isExampleOf",
            "contrastedWith",
            "appliedIn",
            "isPartOf",
            "causeOf",
            "isGeneralizationOf",
        ]:
            return False

        # Check for structural labels (instructional headings, not concepts)
        if self._is_structural_label(triple.subject) or self._is_structural_label(
            triple.object
        ):
            return False

        # Layer 2: Evidence similarity
        threshold = (
            settings.evidence_similarity_threshold_llm
            if is_llm_slide
            else settings.evidence_similarity_threshold
        )
        ev_emb = self.sbert.encode(triple.evidence, convert_to_tensor=True)
        text_emb = self.sbert.encode(sc.clean_text, convert_to_tensor=True)
        sim = util.cos_sim(ev_emb, text_emb).item()

        if sim < threshold:
            return False

        # Layer 4: Semantic post-validation (5 rules)
        # Note: Layer 3 (confidence threshold) removed to match reference implementation
        return self._semantic_validation(triple)

    def _semantic_validation(self, triple: Triple) -> bool:
        """Layer 4 semantic checks (5 rules)."""
        ev_lower = triple.evidence.lower()

        # Stop words for content word extraction
        _STOP = frozenset({"a", "an", "the", "of", "in", "is", "are", "for", "to"})

        def in_ev(phrase: str) -> int:
            """Return start pos of phrase in ev_lower at a word boundary, or -1.

            Sophisticated version matching reference implementation:
            - Tries exact phrase match with word boundaries
            - Tries plural/singular variant
            - For multi-word phrases, tries root content word
            - For root content word, tries singular form if root fails
            """

            def _s(p):
                pat = re.escape(p.lower())
                m = re.search(r"(?<![a-z])" + pat + r"(?![a-z])", ev_lower)
                return m.start() if m else -1

            pos = _s(phrase)
            if pos != -1:
                return pos

            # Try plural/singular variant
            alt = phrase[:-1] if phrase.endswith("s") else phrase + "s"
            pos = _s(alt)
            if pos != -1:
                return pos

            # Try root content word for multi-word phrases only
            # (e.g. "arrays of objects" → "arrays").
            # Do NOT try plural/singular of the root word alone — this causes false matches
            # (e.g. "element" + "s" = "elements" would wrongly match "element type" in evidence).
            # Exception: try the singular of the root word if the root itself fails
            # (e.g. "arrays" → "array") because evidence often uses singular form.
            words = phrase.lower().split()
            content = [w for w in words if w not in _STOP]
            if len(words) > 1 and content and content[0] != phrase.lower():
                pos = _s(content[0])
                if pos != -1:
                    return pos
                # Try singular of root content word only
                root = content[0]
                if root.endswith("s") and len(root) > 3:
                    return _s(root[:-1])
            return -1

        # L4-2: isDefinedAs checks
        if triple.relation == "isDefinedAs":
            signals = [
                " is ",
                " are ",
                " called ",
                " refers to ",
                " means ",
                " defined as ",
                " known as ",
                " denoted ",
                " represents ",
                " represented as ",
                " can be ",
                " consist ",
            ]
            if not any(s in ev_lower for s in signals):
                return False
            if in_ev(triple.subject) == -1 or in_ev(triple.object) == -1:
                return False
            if (
                f"called {triple.object}" in ev_lower
                and f"called {triple.subject}" not in ev_lower
            ):
                return False
            if "can be represented as" in ev_lower:
                return False

        # L4-3: isPartOf usage signals
        if triple.relation == "isPartOf":
            usage_signals = [
                "passed",
                " used in ",
                "operates in",
                " works in ",
                "applied to",
                " referenced ",
                " referenced using ",
                " accessed ",
            ]
            if any(s in ev_lower for s in usage_signals):
                return False

        # L4-4: isExampleOf IS-A check
        if triple.relation == "isExampleOf":
            isa_signals = [
                " is a ",
                " is an ",
                " are a ",
                " such as ",
                " example ",
                " instance ",
                " implementation ",
                " type of ",
                " kind of ",
                " includes ",
                " like ",
                "implementations",
            ]
            has_isa = any(s in ev_lower for s in isa_signals)

            obj_pos = in_ev(triple.object)
            obj_in_ev = obj_pos != -1
            obj_injected = any(
                k.phrase == triple.object and k.source_type == "injected" for k in []
            )

            if not has_isa:
                if not obj_in_ev and not obj_injected:
                    return False
                if not obj_in_ev and obj_injected:
                    if in_ev(triple.subject) == -1:
                        return False
                if obj_in_ev:
                    subj_pos = in_ev(triple.subject)
                    if subj_pos != -1 and subj_pos > obj_pos:
                        return False

            if "making" in ev_lower:
                subj_pos = in_ev(triple.subject)
                obj_pos_check = in_ev(triple.object)
                making_pos = ev_lower.find("making")
                if subj_pos != -1 and obj_pos_check != -1:
                    if subj_pos < making_pos < obj_pos_check:
                        return False
                making_pos = ev_lower.find("making")
                if subj_pos < making_pos < obj_pos:
                    return False

        # L4-5: isGeneralizationOf helper signals
        if triple.relation == "isGeneralizationOf":
            helper_signals = [
                " helper ",
                " abstract class for ",
                " utility class ",
                " base class for ",
                " supports ",
                " assists ",
            ]
            if any(s in ev_lower for s in helper_signals):
                return False

        return True

    def _phrases_are_synonyms(self, a: str, b: str) -> bool:
        """
        Return True when two keyphrase strings refer to the same concept and
        should be treated as identical for deduplication purposes.

        Two cases handled — both are generic (no hardcoding):

        1. Acronym / expansion pairs:
           One string is an acronym (short all-alpha, ≤6 chars) and the other
           is its expansion — the acronym matches the first letter of EVERY word
           in the expansion (no stop-word filtering, to handle cases like
           "first in, first out" → f,i,f,o → "fifo").
           e.g. "fifo" vs "first in, first out"
                "lifo" vs "last in, first out"
                "api"  vs "application programming interface"

        2. Substring containment with word-boundary guard:
           One phrase is a clean substring of the other AND the match ends at a
           word boundary (space or end of string) to avoid false positives like
           "array" matching inside "arraylist".
           e.g. "ordered list" vs "ordered list of values" ✓
                "array" vs "arraylist"                     ✗ (no word boundary)
           Capped at 5 words for the shorter phrase.
        """
        a, b = a.lower().strip(), b.lower().strip()
        if a == b:
            return True

        # ── Case 1: acronym / expansion ───────────────────────────────────────────
        def _is_acronym_of(short: str, long: str) -> bool:
            """Return True if `short` is the acronym formed from initials of `long`."""
            short_clean = re.sub(r"[^a-z]", "", short)
            if len(short_clean) < 2 or len(short_clean) > 6:
                return False
            # Use ALL word initials (no stop-word filtering) so that
            # "first in, first out" → [f,i,f,o] → "fifo" matches correctly
            long_words = re.sub(r"[^a-z\s]", "", long).split()
            if not long_words:
                return False
            all_initials = "".join(w[0] for w in long_words if w and w[0].isalpha())
            return short_clean == all_initials

        if _is_acronym_of(a, b) or _is_acronym_of(b, a):
            return True

        # ── Case 2: substring containment with word-boundary guard ────────────────
        shorter, longer = (a, b) if len(a) <= len(b) else (b, a)
        shorter_words = shorter.split()
        if len(shorter_words) < 1 or len(shorter_words) > 5:
            return False
        idx = longer.find(shorter)
        if idx == -1:
            return False
        # Word-boundary check: char immediately after the match must be a space or EOL
        after = longer[idx + len(shorter) :]
        if after and after[0] != " ":
            return False  # e.g. "array" inside "arraylist" — next char is "l", not " "
        # Also require they share the same first word
        if not longer.startswith(shorter_words[0]):
            return False
        return True

    def _dedup_triples(self, triples: List[Triple]) -> List[Triple]:
        """
        Deduplicate triples within the same slide.

        Two deduplication passes:

        Pass 1 — Exact pair dedup:
          For any two triples with the same (subject, object) pair, keep only
          the one with the highest confidence. Handles cases where the LLM
          generates both isPartOf and isExampleOf for the same pair.

        Pass 2 — Semantic pair dedup:
          For any two triples where the subjects are synonyms AND the objects
          are synonyms (or vice versa), keep only the higher-confidence one.
          This handles acronym/expansion duplicates like:
            "queue isDefinedAs fifo" + "queue isDefinedAs first in, first out"
          where the two objects mean the same thing but are different strings,
          so Pass 1 would not catch them.

        Both passes are fully generic — no domain-specific knowledge used.
        """
        if not triples:
            return []

        # Pass 1: exact pair dedup
        seen_pairs: Dict[Tuple[str, str], Triple] = {}
        for t in triples:
            pair = (t.subject, t.object)
            if pair not in seen_pairs or t.confidence > seen_pairs[pair].confidence:
                seen_pairs[pair] = t

        triples = list(seen_pairs.values())

        # Pass 2: semantic pair dedup — O(n²) but n is small per slide (≤ 10 triples)
        final: List[Triple] = []
        for t in sorted(triples, key=lambda x: -x.confidence):
            is_semantic_dup = False
            for existing in final:
                # Same relation, semantically equivalent subject AND object
                subj_match = self._phrases_are_synonyms(t.subject, existing.subject)
                obj_match = self._phrases_are_synonyms(t.object, existing.object)
                # OR reversed direction (both encode same real-world relationship)
                subj_obj_match = self._phrases_are_synonyms(t.subject, existing.object)
                obj_subj_match = self._phrases_are_synonyms(t.object, existing.subject)
                if (subj_match and obj_match) or (subj_obj_match and obj_subj_match):
                    is_semantic_dup = True
                    break
            if not is_semantic_dup:
                final.append(t)

        return final

    def _phrases_are_synonyms(self, a: str, b: str) -> bool:
        """Check if two phrases are near-synonyms."""
        # Acronym check
        a_alpha = "".join(c for c in a if c.isalpha())
        b_alpha = "".join(c for c in b if c.isalpha())
        if len(a_alpha) <= 6 and len(b_alpha) > len(a_alpha):
            if all(c in b_alpha.lower() for c in a_alpha.lower()):
                return True
        if len(b_alpha) <= 6 and len(a_alpha) > len(b_alpha):
            if all(c in a_alpha.lower() for c in b_alpha.lower()):
                return True

        # Substring containment
        if a in b or b in a:
            # Word boundary check
            if a in b:
                idx = b.find(a)
                if (idx == 0 or not b[idx - 1].isalnum()) and (
                    idx + len(a) == len(b) or not b[idx + len(a)].isalnum()
                ):
                    return True
            if b in a:
                idx = a.find(b)
                if (idx == 0 or not a[idx - 1].isalnum()) and (
                    idx + len(b) == len(a) or not a[idx + len(b)].isalnum()
                ):
                    return True

        return False

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 5-9 continue (simplified for space - full implementation needed)
    # ══════════════════════════════════════════════════════════════════════════

    def step5_weight_and_prune(
        self,
        triples: List[Triple],
        keyphrases: Dict[str, List[Keyphrase]],
        slides: List[SlideContent],
    ) -> Tuple[List[ConceptNode], List[Triple]]:
        """
        Weight concepts using 3-signal SBERT and prune low-weight nodes.

        Weight formula:
        final_weight = (0.5 × w_evidence) + (0.3 × w_slide) + (0.2 × w_doc)
                       + relation_role_boost + source_type_boost
        """
        from core.config import RELATION_ROLE_BOOSTS, SOURCE_TYPE_BOOSTS

        print(f"[Step 5] Weighting and pruning concepts")
        print(f"  Slides loaded          : {len(slides)}")
        print(f"  Slides with keyphrases : {len(keyphrases)}")
        print(f"  Triples loaded         : {len(triples)}\n")

        # ── Build lookup structures ───────────────────────────────────────────────
        slide_text_map: Dict[str, str] = {s.slide_id: s.clean_text for s in slides}
        full_doc_text = " ".join(s.clean_text for s in slides if s.clean_text.strip())

        print("Encoding full document ...")
        doc_emb = self.sbert.encode(
            full_doc_text, convert_to_tensor=True, show_progress_bar=False
        )
        print("Done.\n")

        # ── Collect unique concepts ───────────────────────────────────────────────
        concept_source_map: Dict[str, dict] = {}

        for slide_id, kps in keyphrases.items():
            for kp in kps:
                phrase = kp.phrase.lower().strip()
                if phrase not in concept_source_map:
                    concept_source_map[phrase] = {
                        "source_type": kp.source_type,
                        "keyphrase_score": kp.score,
                        "slide_ids": [],
                    }
                if slide_id not in concept_source_map[phrase]["slide_ids"]:
                    concept_source_map[phrase]["slide_ids"].append(slide_id)

        for triple in triples:
            for role in ("subject", "object"):
                phrase = (
                    triple.subject.lower().strip()
                    if role == "subject"
                    else triple.object.lower().strip()
                )
                if phrase not in concept_source_map:
                    concept_source_map[phrase] = {
                        "source_type": "body",
                        "keyphrase_score": 0.0,
                        "slide_ids": [],
                    }
                if triple.slide_id not in concept_source_map[phrase]["slide_ids"]:
                    concept_source_map[phrase]["slide_ids"].append(triple.slide_id)

        print(f"Total unique concepts to weight: {len(concept_source_map)}")

        # ── Relation role boosts ──────────────────────────────────────────────────
        concept_role_boost: Dict[str, float] = {p: 0.0 for p in concept_source_map}
        for triple in triples:
            subj = triple.subject.lower().strip()
            obj = triple.object.lower().strip()
            rel = triple.relation
            concept_role_boost[subj] += RELATION_ROLE_BOOSTS.get((rel, "subject"), 0.0)
            concept_role_boost[obj] += RELATION_ROLE_BOOSTS.get((rel, "object"), 0.0)

        # ── Evidence sentence map ─────────────────────────────────────────────────
        concept_evidence_map: Dict[str, List[str]] = {p: [] for p in concept_source_map}
        for triple in triples:
            if triple.evidence:
                concept_evidence_map[triple.subject.lower().strip()].append(
                    triple.evidence
                )
                concept_evidence_map[triple.object.lower().strip()].append(
                    triple.evidence
                )

        # ── Weight computation ────────────────────────────────────────────────────
        def compute_weight(phrase: str) -> float:
            info = concept_source_map[phrase]
            slide_ids = info["slide_ids"]
            src_type = info["source_type"]

            phrase_emb = self.sbert.encode(
                phrase, convert_to_tensor=True, show_progress_bar=False
            )

            evidence_sents = concept_evidence_map.get(phrase, [])
            if evidence_sents:
                best_ev_sim = max(
                    float(
                        util.cos_sim(
                            phrase_emb,
                            self.sbert.encode(
                                ev, convert_to_tensor=True, show_progress_bar=False
                            ),
                        )
                    )
                    for ev in evidence_sents
                )
            else:
                best_ev_sim = 0.0

            if slide_ids:
                texts = [
                    slide_text_map.get(sid, "")
                    for sid in slide_ids
                    if slide_text_map.get(sid)
                ]
                if texts:
                    slide_emb = self.sbert.encode(
                        " ".join(texts), convert_to_tensor=True, show_progress_bar=False
                    )
                    w_slide = float(util.cos_sim(phrase_emb, slide_emb))
                else:
                    w_slide = 0.0
            else:
                w_slide = 0.0

            w_doc = float(util.cos_sim(phrase_emb, doc_emb))

            raw = (
                settings.w_evidence * best_ev_sim
                + settings.w_slide * w_slide
                + settings.w_doc * w_doc
            )
            role_boost = min(concept_role_boost.get(phrase, 0.0), 0.20)
            src_boost = SOURCE_TYPE_BOOSTS.get(src_type, 0.0)
            return round(min(raw + role_boost + src_boost, 1.0), 4)

        # ── Run weighting ─────────────────────────────────────────────────────────
        print("\nComputing weights ...")
        concept_weights: Dict[str, float] = {}
        for i, phrase in enumerate(concept_source_map, 1):
            concept_weights[phrase] = compute_weight(phrase)
            if i % 10 == 0 or i == len(concept_source_map):
                print(f"  [{i}/{len(concept_source_map)}] done")

        # ── Semantic merge ────────────────────────────────────────────────────────
        print("\nRunning semantic conflict resolution ...")
        phrases_sorted = sorted(concept_weights, key=lambda p: -concept_weights[p])
        merged_into: Dict[str, str] = {}
        needs_review: set = set()

        phrase_embs = {
            p: self.sbert.encode(p, convert_to_tensor=True, show_progress_bar=False)
            for p in phrases_sorted
        }

        for i, phrase in enumerate(phrases_sorted):
            if phrase in merged_into:
                continue
            for other in phrases_sorted[i + 1 :]:
                if other in merged_into:
                    continue
                sim = float(util.cos_sim(phrase_embs[phrase], phrase_embs[other]))
                if sim >= settings.merge_sim_threshold:
                    merged_into[other] = phrase
                    print(f"  MERGE  '{other}' → '{phrase}'  (sim={sim:.3f})")
                elif sim >= settings.review_sim_threshold:
                    needs_review.add(phrase)
                    needs_review.add(other)

        print(
            f"  Merged : {len(merged_into)}  |  Flagged for review : {len(needs_review)}"
        )

        # ── Build ConceptNode objects ─────────────────────────────────────────────
        concept_nodes: Dict[str, ConceptNode] = {}

        for phrase, info in concept_source_map.items():
            canonical = merged_into.get(phrase, phrase)
            if canonical != phrase:
                if (
                    canonical in concept_nodes
                    and phrase not in concept_nodes[canonical].aliases
                ):
                    concept_nodes[canonical].aliases.append(phrase)
                continue

            aliases = [o for o, c in merged_into.items() if c == phrase]
            all_slide_ids = list(info["slide_ids"])
            for alias in aliases:
                for sid in concept_source_map.get(alias, {}).get("slide_ids", []):
                    if sid not in all_slide_ids:
                        all_slide_ids.append(sid)

            concept_nodes[phrase] = ConceptNode(
                name=phrase,
                aliases=aliases,
                slide_ids=sorted(all_slide_ids),
                source_type=info["source_type"],
                keyphrase_score=info["keyphrase_score"],
                final_weight=concept_weights[phrase],
                doc_id=self.doc_id,
                needs_review=phrase in needs_review,
            )

        # ── Prune ─────────────────────────────────────────────────────────────────
        before = len(concept_nodes)
        pruned_names = {
            n
            for n, node in concept_nodes.items()
            if node.final_weight < settings.weight_threshold
        }
        concept_nodes = {
            k: v for k, v in concept_nodes.items() if k not in pruned_names
        }
        after = len(concept_nodes)

        surviving_names = set(concept_nodes.keys())
        for node in concept_nodes.values():
            surviving_names.update(node.aliases)

        triples_before = len(triples)
        kept_triples = [
            t
            for t in triples
            if t.subject.lower().strip() in surviving_names
            and t.object.lower().strip() in surviving_names
        ]

        # ── Summary ───────────────────────────────────────────────────────────────
        print(f"\n{'=' * 60}")
        print(f"Step 5 complete  [{self.doc_id}]")
        print(f"  Concepts  : {before} → {after}  (pruned {before - after})")
        print(f"  Triples   : {triples_before} → {len(kept_triples)}")
        print(
            f"  Merged    : {len(merged_into)}  |  Review flags : {len(needs_review)}"
        )
        print(f"\nTop 15 concepts by weight:")
        for node in sorted(concept_nodes.values(), key=lambda n: -n.final_weight)[:15]:
            review_flag = " ⚑" if node.needs_review else ""
            print(
                f"  {node.name:<35} {node.final_weight:.4f}  [{node.source_type}]{review_flag}"
            )

        # ── Save ──────────────────────────────────────────────────────────────────
        concepts_list = list(concept_nodes.values())
        self._save_json([asdict(c) for c in concepts_list], "step5_concepts")
        self._save_json([asdict(t) for t in kept_triples], "step5_triples_pruned")

        return concepts_list, kept_triples

    def step6_expand(
        self,
        concepts: List[ConceptNode],
        triples: List[Triple],
        slides: List[SlideContent],
        keyphrases: Dict[str, List[Keyphrase]],
    ) -> List[ExpansionEdge]:
        """
        Closed-corpus concept expansion.

        For each core concept (subject/object in a validated triple), asks the LLM
        to select related concepts from the document vocabulary.
        Every candidate is gated by SBERT (>= 0.65) then a slide-scope filter.
        Produces "relatedConcept" expansion edges.

        4-phase expansion per core concept:
        1. LLM selection — ask LLM to select related concepts from pool
        2. SBERT gate — keep only candidates with cosine ≥ 0.65 vs core concept
        3. Slide-scope constraint — keep only candidates that are:
           (a) on an adjacent slide (±1 index), OR
           (b) have SBERT cosine ≥ 0.70 vs full document
        4. Deduplication — skip self-loops, trivial pairs, already-existing pairs
        """
        print(f"[Step 6] Expanding concepts")

        # ── Build document vocabulary ─────────────────────────────────────────────
        print("\nPhase 1: Building document vocabulary ...")
        concept_slide_map: Dict[str, set] = {}

        for slide_id, kp_list in keyphrases.items():
            for kp in kp_list:
                phrase = kp.phrase.lower().strip()
                concept_slide_map.setdefault(phrase, set()).add(slide_id)

        for t in triples:
            for role in ("subject", "object"):
                phrase = (
                    t.subject.lower().strip()
                    if role == "subject"
                    else t.object.lower().strip()
                )
                concept_slide_map.setdefault(phrase, set()).add(t.slide_id)

        # Augment with spaCy noun chunks
        try:
            nlp = spacy.load("en_core_web_sm")
            print("  Adding spaCy noun chunks ...")
            for s in slides:
                if not s.clean_text.strip():
                    continue
                doc_nlp = nlp(s.clean_text)
                for chunk in doc_nlp.noun_chunks:
                    phrase = chunk.text.lower().strip()
                    stripped = re.sub(r"^(a|an|the)\s+", "", phrase).strip()
                    if stripped in concept_slide_map:
                        continue
                    if len(phrase) >= 3 and not all(tok.is_stop for tok in chunk):
                        concept_slide_map.setdefault(phrase, set()).add(s.slide_id)
        except Exception as e:
            print(f"  spaCy noun chunks skipped: {e}")

        doc_vocabulary: List[str] = sorted(concept_slide_map.keys())
        print(f"  Vocabulary size: {len(doc_vocabulary)} concepts")

        # ── Pre-encode vocabulary ─────────────────────────────────────────────────
        print("\nPre-encoding vocabulary ...")
        vocab_embs: Dict[str, any] = {
            phrase: self.sbert.encode(
                phrase, convert_to_tensor=True, show_progress_bar=False
            )
            for phrase in doc_vocabulary
        }
        print(f"  Encoded {len(vocab_embs)} embeddings.")

        # ── Encode full document ──────────────────────────────────────────────────
        slide_text_map: Dict[str, str] = {s.slide_id: s.clean_text for s in slides}
        full_doc_text = " ".join(s.clean_text for s in slides if s.clean_text.strip())
        doc_emb = self.sbert.encode(
            full_doc_text, convert_to_tensor=True, show_progress_bar=False
        )

        # ── Ordered slide index (for adjacency check) ─────────────────────────────
        ordered_slides: List[str] = sorted(slide_text_map.keys())
        slide_index: Dict[str, int] = {sid: i for i, sid in enumerate(ordered_slides)}

        # ── Existing edge pairs (avoid duplicating extraction edges) ──────────────
        existing_pairs: set = set()
        for t in triples:
            s_n = t.subject.lower().strip()
            o_n = t.object.lower().strip()
            existing_pairs.add((s_n, o_n))
            existing_pairs.add((o_n, s_n))

        # ── LLM prompts ───────────────────────────────────────────────────────────
        S6_SYSTEM = (
            "You are a knowledge graph assistant for educational materials.\n"
            "Given a main concept and its slide context, select the most educationally "
            "relevant related concepts from the candidate pool.\n\n"
            "STRICT RULES:\n"
            "1. Select ONLY concepts from the candidate pool. Do not invent anything.\n"
            "2. Choose concepts meaningfully related in an educational sense.\n"
            f"3. Select at most {settings.max_candidates_per_concept} concepts.\n"
            '4. Return ONLY valid JSON: {{"related": ["concept1", "concept2"]}}\n'
            '5. If nothing is relevant return: {{"related": []}}\n'
            "6. Do not explain your choices."
        )

        S6_USER = (
            "Main concept: '{concept}'\n"
            "Slide context: '{slide_context}'\n\n"
            "Candidate pool (select ONLY from this list):\n"
            "{pool}\n\n"
            'Return JSON with key "related" containing an array of selected strings.'
        )

        # ── Initialize LLM ────────────────────────────────────────────────────────
        if self.llm_primary is None:
            self.llm_primary = ChatGroq(
                model=settings.llm_primary,
                temperature=0,
                api_key=settings.groq_api_key,
                model_kwargs={"response_format": {"type": "json_object"}},
            )
        if self.llm_fallback is None:
            self.llm_fallback = ChatGroq(
                model=settings.llm_fallback,
                temperature=0,
                api_key=settings.groq_api_key,
                model_kwargs={"response_format": {"type": "json_object"}},
            )

        # ── Helper: invoke LLM with fallback ──────────────────────────────────────
        def invoke_llm(messages: list, max_retries: int = 3) -> Optional[str]:
            for attempt in range(1, max_retries + 1):
                try:
                    return str(self.llm_primary.invoke(messages).content)
                except Exception as e:
                    if "429" in str(e) or "rate" in str(e).lower():
                        print(
                            f"    Rate-limited (attempt {attempt}) — trying fallback ..."
                        )
                        try:
                            return str(self.llm_fallback.invoke(messages).content)
                        except Exception as e2:
                            if "429" in str(e2) or "rate" in str(e2).lower():
                                wait = 5 * attempt
                                print(f"    Still rate-limited — waiting {wait}s ...")
                                time.sleep(wait)
                                continue
                            print(f"    Fallback error: {e2}")
                            return None
                    print(f"    LLM error: {e}")
                    return None
            return None

        # ── Run expansion ─────────────────────────────────────────────────────────
        print("\nPhases 2–4: LLM selection + SBERT gate + slide-scope filter ...")
        core_concepts = sorted(
            set(
                role_val.lower().strip()
                for t in triples
                for role_val in [t.subject, t.object]
            )
        )
        print(f"Core concepts to expand: {len(core_concepts)}\n")

        expansion_edges: List[ExpansionEdge] = []
        s6_stats = {
            "llm_calls": 0,
            "proposed": 0,
            "sbert_pass": 0,
            "scope_pass": 0,
            "added": 0,
        }

        for i, concept in enumerate(core_concepts, 1):
            concept_slides = sorted(concept_slide_map.get(concept, set()))
            if not concept_slides:
                print(f"  [{i:2d}/{len(core_concepts)}] {concept}: no slide — skipped")
                continue

            primary_slide = concept_slides[0]
            slide_context = slide_text_map.get(primary_slide, "")[:400]

            already_linked = {o for (s_n, o) in existing_pairs if s_n == concept} | {
                s_n for (s_n, o) in existing_pairs if o == concept
            }
            pool = [
                p
                for p in doc_vocabulary
                if p != concept and p not in already_linked and len(p) >= 3
            ]
            if not pool:
                continue

            messages = [
                SystemMessage(content=S6_SYSTEM),
                HumanMessage(
                    content=S6_USER.format(
                        concept=concept,
                        slide_context=slide_context,
                        pool="\n".join(f"- {p}" for p in pool[:80]),
                    )
                ),
            ]

            raw = invoke_llm(messages)
            s6_stats["llm_calls"] += 1

            try:
                candidates = [
                    str(c).lower().strip()
                    for c in json.loads(raw or "{}").get("related", [])
                ]
            except Exception:
                candidates = []

            candidates = [c for c in candidates if c in vocab_embs]
            s6_stats["proposed"] += len(candidates)

            if not candidates:
                print(f"  [{i:2d}/{len(core_concepts)}] {concept:<35} 0 candidates")
                continue

            # Phase 3: SBERT gate
            c_emb = (
                vocab_embs[concept]
                if concept in vocab_embs
                else self.sbert.encode(concept, convert_to_tensor=True)
            )
            sbert_passed = [
                (cand, round(float(util.cos_sim(c_emb, vocab_embs[cand])), 4))
                for cand in candidates
                if float(util.cos_sim(c_emb, vocab_embs[cand]))
                >= settings.sbert_sim_threshold
            ]
            s6_stats["sbert_pass"] += len(sbert_passed)

            # Phase 4: slide-scope filter
            c_indices = {slide_index[s] for s in concept_slides if s in slide_index}
            scope_passed = []
            for cand, sim in sbert_passed:
                cand_slides = concept_slide_map.get(cand, set())
                cand_indices = {slide_index[s] for s in cand_slides if s in slide_index}
                adjacent = any(
                    abs(ci - cj) <= 1 for ci in c_indices for cj in cand_indices
                )
                high_doc_rel = (
                    float(util.cos_sim(vocab_embs[cand], doc_emb))
                    >= settings.doc_weight_threshold
                )
                if adjacent or high_doc_rel:
                    scope_passed.append((cand, sim))
            s6_stats["scope_pass"] += len(scope_passed)

            # Add edges
            added = 0
            for cand, sim in scope_passed:
                if cand == concept:
                    continue
                if re.sub(r"^(a|an|the)\s+", "", cand).strip() == concept:
                    continue
                pair = (concept, cand)
                if pair in existing_pairs:
                    continue
                existing_pairs.add(pair)
                existing_pairs.add((cand, concept))
                expansion_edges.append(
                    ExpansionEdge(
                        subject=concept,
                        relation="relatedConcept",
                        object=cand,
                        source="expansion",
                        confidence=sim,
                        slide_id=primary_slide,
                        doc_id=self.doc_id,
                    )
                )
                added += 1

            s6_stats["added"] += added
            print(
                f"  [{i:2d}/{len(core_concepts)}] {concept:<35} "
                f"LLM={len(candidates)}  SBERT={len(sbert_passed)}  "
                f"scope={len(scope_passed)}  added={added}"
            )

        # ── Summary ───────────────────────────────────────────────────────────────
        print(f"\n{'=' * 60}")
        print(f"Step 6 complete  [{self.doc_id}]")
        print(f"  Core concepts processed : {len(core_concepts)}")
        print(f"  LLM calls               : {s6_stats['llm_calls']}")
        print(f"  Candidates proposed     : {s6_stats['proposed']}")
        print(f"  Passed SBERT gate       : {s6_stats['sbert_pass']}")
        print(f"  Passed scope filter     : {s6_stats['scope_pass']}")
        print(f"  Expansion edges added   : {s6_stats['added']}")

        if expansion_edges:
            print("\nSample expansion edges (first 5):")
            for e in expansion_edges[:5]:
                print(
                    f"  '{e.subject}' --[relatedConcept]--> '{e.object}'  conf={e.confidence:.3f}"
                )

        # ── Save ──────────────────────────────────────────────────────────────────
        self._save_json([asdict(e) for e in expansion_edges], "step6_expansion")

        return expansion_edges

    def step7_store_neo4j(
        self,
        concepts: List[ConceptNode],
        triples: List[Triple],
        expansions: List[ExpansionEdge],
        slides: List[SlideContent],
    ):
        """
        Store graph in Neo4j AuraDB.

        Three-pass write strategy:
        1. Store all Concept nodes first (nodes must exist before edges)
        2. Store edges slide-by-slide (extraction + expansion)
        3. Create LearningMaterial node + BELONGS_TO links

        Uses MERGE (never CREATE) to avoid duplicates.
        Applies semantic conflict resolution at write time (0.92 threshold).
        """
        from collections import defaultdict

        print(f"[Step 7] Storing in Neo4j")

        print(f"Concepts           : {len(concepts)}")
        print(f"Extraction triples : {len(triples)}")
        print(f"Expansion edges    : {len(expansions)}")

        # ── Build concept lookup ──────────────────────────────────────────────────
        concept_by_name: Dict[str, ConceptNode] = {}
        for c in concepts:
            concept_by_name[c.name] = c
            for alias in c.aliases:
                concept_by_name[alias] = c

        # ── Pre-encode for conflict resolution ────────────────────────────────────
        print("\nPre-encoding concept names ...")
        concept_embs_s7 = {
            c.name: self.sbert.encode(
                c.name, convert_to_tensor=True, show_progress_bar=False
            )
            for c in concepts
        }

        def resolve_name(name: str) -> str:
            """Resolve name to canonical form using semantic similarity."""
            nl = name.lower().strip()
            if nl in concept_by_name:
                return concept_by_name[nl].name
            emb = self.sbert.encode(nl, convert_to_tensor=True, show_progress_bar=False)
            for cname, cemb in concept_embs_s7.items():
                if float(util.cos_sim(emb, cemb)) >= settings.merge_sim_threshold:
                    return cname
            return nl

        # ── Connect ───────────────────────────────────────────────────────────────
        print(f"\nConnecting to Neo4j at {settings.neo4j_uri} ...")
        driver = GraphDatabase.driver(
            settings.neo4j_uri, auth=(settings.neo4j_user, settings.neo4j_password)
        )

        def run_q(tx, query, **params):
            return tx.run(query, **params)

        # Create constraints and indexes
        with driver.session() as session:
            session.run(
                "CREATE CONSTRAINT concept_name IF NOT EXISTS FOR (c:Concept) REQUIRE c.name IS UNIQUE"
            )
            session.run(
                "CREATE INDEX concept_doc IF NOT EXISTS FOR (c:Concept) ON (c.doc_id)"
            )
            session.run(
                "CREATE INDEX lm_doc IF NOT EXISTS FOR (l:LearningMaterial) ON (l.doc_id)"
            )
        print("Constraints and indexes ready.")

        # ── Cypher queries ────────────────────────────────────────────────────────
        UPSERT_CONCEPT = """
        MERGE (c:Concept {name: $name})
        SET c.aliases=$aliases, c.slide_ids=$slide_ids, c.source_type=$source_type,
            c.keyphrase_score=$keyphrase_score, c.final_weight=$final_weight,
            c.doc_id=$doc_id, c.needs_review=$needs_review
        """

        CREATE_EXTR_EDGE = """
        MATCH (s:Concept {name:$subject}),(o:Concept {name:$object})
        MERGE (s)-[r:RELATION {relation_type:$relation_type,slide_id:$slide_id,doc_id:$doc_id}]->(o)
        SET r.evidence=$evidence, r.confidence=$confidence, r.source='extraction'
        """

        CREATE_EXP_EDGE = """
        MATCH (s:Concept {name:$subject}),(o:Concept {name:$object})
        MERGE (s)-[r:RELATION {relation_type:'relatedConcept',slide_id:$slide_id,doc_id:$doc_id}]->(o)
        SET r.confidence=$confidence, r.source='expansion'
        """

        UPSERT_LM = """
        MERGE (lm:LearningMaterial {doc_id:$doc_id})
        SET lm.title=$title, lm.total_slides=$total_slides
        """

        LINK_LM = """
        MATCH (c:Concept {name:$name,doc_id:$doc_id}),(lm:LearningMaterial {doc_id:$doc_id})
        MERGE (c)-[r:BELONGS_TO]->(lm)
        SET r.final_weight=$final_weight
        """

        # ── Group by slide ────────────────────────────────────────────────────────
        slide_text_map = {s.slide_id: s.clean_text for s in slides}
        triples_by_slide: Dict[str, List[Triple]] = defaultdict(list)
        expansion_by_slide: Dict[str, List[ExpansionEdge]] = defaultdict(list)

        for t in triples:
            triples_by_slide[t.slide_id].append(t)
        for e in expansions:
            expansion_by_slide[e.slide_id].append(e)

        ordered_slides_s7 = sorted(slide_text_map.keys())

        report = {
            "doc_id": self.doc_id,
            "total_slides": len(ordered_slides_s7),
            "slides_stored": [],
            "total_concepts_stored": 0,
            "total_extraction_edges": 0,
            "total_expansion_edges": 0,
            "errors": [],
        }

        # ── Pass 1: Concept nodes ─────────────────────────────────────────────────
        print("\nPass 1: Storing concept nodes ...")
        nodes_stored = 0
        with driver.session() as session:
            for c in concepts:
                try:
                    session.execute_write(
                        run_q,
                        UPSERT_CONCEPT,
                        name=c.name,
                        aliases=c.aliases,
                        slide_ids=c.slide_ids,
                        source_type=c.source_type,
                        keyphrase_score=float(c.keyphrase_score),
                        final_weight=float(c.final_weight),
                        doc_id=c.doc_id,
                        needs_review=bool(c.needs_review),
                    )
                    nodes_stored += 1
                except Exception as e:
                    report["errors"].append(f"node '{c.name}': {e}")

            # Ensure expansion endpoints exist
            exp_names = {
                role_val.lower().strip()
                for e in expansions
                for role_val in [e.subject, e.object]
            }
            for name in exp_names:
                if name not in concept_by_name:
                    try:
                        session.execute_write(
                            run_q,
                            UPSERT_CONCEPT,
                            name=name,
                            aliases=[],
                            slide_ids=[],
                            source_type="body",
                            keyphrase_score=0.0,
                            final_weight=0.0,
                            doc_id=self.doc_id,
                            needs_review=False,
                        )
                        nodes_stored += 1
                    except Exception:
                        pass

        print(f"  Stored {nodes_stored} concept nodes.")

        # ── Pass 2: Edges slide-by-slide ─────────────────────────────────────────
        print("\nPass 2: Storing edges slide by slide ...")
        for slide_id in ordered_slides_s7:
            sr = {
                "slide_id": slide_id,
                "extraction_edges": 0,
                "expansion_edges": 0,
                "errors": [],
            }

            with driver.session() as session:
                # Extraction edges
                for triple in triples_by_slide.get(slide_id, []):
                    subj = resolve_name(triple.subject)
                    obj = resolve_name(triple.object)
                    try:
                        session.execute_write(
                            run_q,
                            CREATE_EXTR_EDGE,
                            subject=subj,
                            object=obj,
                            relation_type=triple.relation,
                            evidence=triple.evidence,
                            confidence=float(triple.confidence),
                            slide_id=slide_id,
                            doc_id=self.doc_id,
                        )
                        sr["extraction_edges"] += 1
                        report["total_extraction_edges"] += 1
                    except Exception as e:
                        msg = f"{slide_id} extr ({subj}->{obj}): {e}"
                        sr["errors"].append(msg)
                        report["errors"].append(msg)

                # Expansion edges
                for edge in expansion_by_slide.get(slide_id, []):
                    subj = resolve_name(edge.subject)
                    obj = resolve_name(edge.object)
                    try:
                        session.execute_write(
                            run_q,
                            CREATE_EXP_EDGE,
                            subject=subj,
                            object=obj,
                            confidence=float(edge.confidence),
                            slide_id=slide_id,
                            doc_id=self.doc_id,
                        )
                        sr["expansion_edges"] += 1
                        report["total_expansion_edges"] += 1
                    except Exception as e:
                        msg = f"{slide_id} exp ({subj}->{obj}): {e}"
                        sr["errors"].append(msg)
                        report["errors"].append(msg)

            status = "OK" if not sr["errors"] else f"ERRORS:{len(sr['errors'])}"
            print(
                f"  {slide_id}: extr={sr['extraction_edges']}  exp={sr['expansion_edges']}  [{status}]"
            )
            report["slides_stored"].append(sr)

        # ── Pass 3: LearningMaterial ──────────────────────────────────────────────
        print("\nPass 3: Creating LearningMaterial node ...")
        with driver.session() as session:
            session.execute_write(
                run_q,
                UPSERT_LM,
                doc_id=self.doc_id,
                title=self.doc_id,
                total_slides=len(ordered_slides_s7),
            )
            for c in concepts:
                try:
                    session.execute_write(
                        run_q,
                        LINK_LM,
                        name=c.name,
                        doc_id=self.doc_id,
                        final_weight=float(c.final_weight),
                    )
                except Exception as e:
                    report["errors"].append(f"BELONGS_TO '{c.name}': {e}")

        # ── Verification ──────────────────────────────────────────────────────────
        print("\nVerification:")
        with driver.session() as session:
            n_concepts = session.run(
                "MATCH (c:Concept {doc_id:$d}) RETURN count(c) AS n", d=self.doc_id
            ).single()["n"]
            n_edges = session.run(
                "MATCH ()-[r:RELATION {doc_id:$d}]->() RETURN count(r) AS n",
                d=self.doc_id,
            ).single()["n"]
            n_extr = session.run(
                "MATCH ()-[r:RELATION {doc_id:$d,source:'extraction'}]->() RETURN count(r) AS n",
                d=self.doc_id,
            ).single()["n"]
            n_exp = session.run(
                "MATCH ()-[r:RELATION {doc_id:$d,source:'expansion'}]->() RETURN count(r) AS n",
                d=self.doc_id,
            ).single()["n"]
            rel_dist = session.run(
                "MATCH ()-[r:RELATION {doc_id:$d}]->() RETURN r.relation_type AS rel, count(r) AS n ORDER BY n DESC",
                d=self.doc_id,
            ).data()

        report["total_concepts_stored"] = n_concepts

        print(f"  Concept nodes  : {n_concepts}")
        print(f"  Total edges    : {n_edges}  (extraction={n_extr}, expansion={n_exp})")
        print(f"  Relations      :")
        for row in rel_dist:
            print(f"    {row['rel']:<25} {row['n']}")

        if report["errors"]:
            print(f"\n  Errors ({len(report['errors'])}):")
            for err in report["errors"][:10]:
                print(f"    {err}")

        driver.close()

        # ── Save ──────────────────────────────────────────────────────────────────
        self._save_json(report, "step7_storage_report")

        print(f"\nStep 7 complete  [{self.doc_id}]")
        print("Open Neo4j Browser → MATCH (n) RETURN n  to visualise.")

    def step8_aggregate(self, slides: List[SlideContent]) -> dict:
        """
        Aggregate LM-EduKG from Neo4j.

        Verifies 4 PACE-KG merge constraints:
        1. Every slide has ≥1 concept
        2. All local concepts exist in Neo4j
        3. All extraction edges have non-empty evidence
        4. Cross-slide expansion edges are tagged material_level=true

        Exports full graph (nodes + edges) + SRS pool for evaluation.
        """
        from collections import defaultdict

        print(f"[Step 8] Aggregating LM-EduKG")

        # ── Reconnect ─────────────────────────────────────────────────────────────
        print(f"Reconnecting to Neo4j for Step 8 [{self.doc_id}] ...")
        driver = GraphDatabase.driver(
            settings.neo4j_uri, auth=(settings.neo4j_user, settings.neo4j_password)
        )

        def q(query: str, **params):
            with driver.session() as session:
                return session.run(query, **params).data()

        # ── Reload Step 7 report ──────────────────────────────────────────────────
        step7_path = self.output_dir / f"{self.doc_id}_step7_storage_report.json"
        with open(step7_path, encoding="utf-8") as f:
            step7_report: dict = json.load(f)

        # ── Reload intermediate outputs ───────────────────────────────────────────
        concepts_path = self.output_dir / f"{self.doc_id}_step5_concepts.json"
        with open(concepts_path, encoding="utf-8") as f:
            concepts_list = json.load(f)

        triples_path = self.output_dir / f"{self.doc_id}_step5_triples_pruned.json"
        with open(triples_path, encoding="utf-8") as f:
            pruned_triples = json.load(f)

        expansion_path = self.output_dir / f"{self.doc_id}_step6_expansion.json"
        with open(expansion_path, encoding="utf-8") as f:
            expansion_edges_raw = json.load(f)

        ordered_slides_s8 = sorted(s.slide_id for s in slides)
        print(f"Slides: {len(ordered_slides_s8)}\n")

        # ── Constraint 1: slide concept coverage ─────────────────────────────────
        print("Constraint 1: Slide concept coverage ...")
        coverage = q(
            "MATCH (c:Concept {doc_id:$d}) UNWIND c.slide_ids AS sid "
            "RETURN sid, count(c) AS n ORDER BY sid",
            d=self.doc_id,
        )
        coverage_map = {row["sid"]: row["n"] for row in coverage}
        content_sids = {s.slide_id for s in slides if s.clean_text.strip()}
        uncovered = [sid for sid in content_sids if coverage_map.get(sid, 0) == 0]

        if uncovered:
            print(
                f"  WARNING: {len(uncovered)} content slides have no concepts: {uncovered}"
            )
        else:
            print(f"  OK — all content slides covered.")

        # ── Constraint 2: concept completeness ───────────────────────────────────
        print("\nConstraint 2: Concept completeness in Neo4j ...")
        local_names = {c["name"] for c in concepts_list}
        neo4j_names = {
            row["name"]
            for row in q(
                "MATCH (c:Concept {doc_id:$d}) RETURN c.name AS name", d=self.doc_id
            )
        }
        missing = local_names - neo4j_names
        print(
            f"  {'OK — all concepts present.' if not missing else f'WARNING: {len(missing)} missing: {missing}'}"
        )

        # ── Constraint 3: evidence provenance ────────────────────────────────────
        print("\nConstraint 3: Evidence provenance ...")
        n_with_ev = q(
            "MATCH ()-[r:RELATION {doc_id:$d,source:'extraction'}]->() "
            "WHERE r.evidence IS NOT NULL AND r.evidence <> '' RETURN count(r) AS n",
            d=self.doc_id,
        )[0]["n"]
        n_without_ev = q(
            "MATCH ()-[r:RELATION {doc_id:$d,source:'extraction'}]->() "
            "WHERE r.evidence IS NULL OR r.evidence = '' RETURN count(r) AS n",
            d=self.doc_id,
        )[0]["n"]
        print(f"  With evidence    : {n_with_ev}")
        print(
            f"  Without evidence : {n_without_ev}  {'OK' if n_without_ev == 0 else 'WARNING'}"
        )

        # ── Constraint 4: cross-slide expansion edges ─────────────────────────────
        print("\nConstraint 4: Cross-slide expansion edges ...")
        concept_slide_set: Dict[str, set] = defaultdict(set)
        for c in concepts_list:
            for sid in c.get("slide_ids", []):
                concept_slide_set[c["name"]].add(sid)

        cross_slide = [
            e
            for e in expansion_edges_raw
            if concept_slide_set.get(e["subject"], set()).isdisjoint(
                concept_slide_set.get(e["object"], set())
            )
            and concept_slide_set.get(e["subject"])
            and concept_slide_set.get(e["object"])
        ]

        print(f"  Cross-slide expansion edges: {len(cross_slide)}")
        for e in cross_slide:
            s1 = next(iter(concept_slide_set.get(e["subject"], {"?"})))
            s2 = next(iter(concept_slide_set.get(e["object"], {"?"})))
            print(f"    '{e['subject']}' -> '{e['object']}'  ({s1} → {s2})")

        if cross_slide:
            for e in cross_slide:
                q(
                    "MATCH (s:Concept {name:$sub,doc_id:$d}),(o:Concept {name:$obj,doc_id:$d}) "
                    "MATCH (s)-[r:RELATION {relation_type:'relatedConcept',source:'expansion'}]->(o) "
                    "SET r.material_level = true",
                    sub=e["subject"],
                    obj=e["object"],
                    d=self.doc_id,
                )
            print(f"  Tagged {len(cross_slide)} edges as material_level=true.")

        # ── Update LearningMaterial node ──────────────────────────────────────────
        total_edges_s8 = (
            step7_report["total_extraction_edges"]
            + step7_report["total_expansion_edges"]
        )
        q(
            "MERGE (lm:LearningMaterial {doc_id:$d}) "
            "SET lm.title=$t, lm.total_slides=$ts, lm.total_concepts=$tc, "
            "    lm.total_extraction_edges=$te, lm.total_expansion_edges=$tx, "
            "    lm.total_edges=$tot, lm.pipeline_version='PACE-KG-v1'",
            d=self.doc_id,
            t=self.doc_id,
            ts=step7_report["total_slides"],
            tc=step7_report["total_concepts_stored"],
            te=step7_report["total_extraction_edges"],
            tx=step7_report["total_expansion_edges"],
            tot=total_edges_s8,
        )
        print("\nLearningMaterial node updated.")

        # ── Full graph export ─────────────────────────────────────────────────────
        print("\nExporting full LM-EduKG ...")

        nodes_export = q(
            "MATCH (c:Concept {doc_id:$d}) "
            "RETURN c.name AS name, c.aliases AS aliases, c.slide_ids AS slide_ids, "
            "c.source_type AS source_type, c.keyphrase_score AS keyphrase_score, "
            "c.final_weight AS final_weight, c.needs_review AS needs_review "
            "ORDER BY c.final_weight DESC",
            d=self.doc_id,
        )

        edges_export = q(
            "MATCH (s:Concept {doc_id:$d})-[r:RELATION]->(o:Concept {doc_id:$d}) "
            "RETURN s.name AS subject, r.relation_type AS relation, o.name AS object, "
            "r.evidence AS evidence, r.confidence AS confidence, r.source AS source, "
            "r.slide_id AS slide_id, r.material_level AS material_level "
            "ORDER BY r.confidence DESC",
            d=self.doc_id,
        )

        rel_dist_s8 = q(
            "MATCH ()-[r:RELATION {doc_id:$d}]->() "
            "RETURN r.relation_type AS relation, count(r) AS count ORDER BY count DESC",
            d=self.doc_id,
        )

        # SRS pool (extraction edges only)
        srs_pool = [
            {
                "triple_id": f"t{i + 1:03d}",
                "subject": e["subject"],
                "relation": e["relation"],
                "object": e["object"],
                "evidence": e["evidence"],
                "confidence": e["confidence"],
                "slide_id": e["slide_id"],
                "source": e["source"],
            }
            for i, e in enumerate(edges_export)
            if e["source"] == "extraction"
        ]

        lm_edkg = {
            "doc_id": self.doc_id,
            "pipeline": "PACE-KG-v1",
            "total_slides": step7_report["total_slides"],
            "total_concepts": len(nodes_export),
            "total_edges": len(edges_export),
            "relation_distribution": {
                row["relation"]: row["count"] for row in rel_dist_s8
            },
            "nodes": nodes_export,
            "edges": edges_export,
            "srs_pool": srs_pool,
            "srs_pool_size": len(srs_pool),
        }

        driver.close()

        # ── Final summary ─────────────────────────────────────────────────────────
        print("\n" + "=" * 60)
        print(f"LM-EduKG Final Summary  [{self.doc_id}]")
        print(f"  Total slides       : {step7_report['total_slides']}")
        print(f"  Concept nodes      : {len(nodes_export)}")
        print(f"  Total edges        : {len(edges_export)}")
        print(f"    Extraction       : {step7_report['total_extraction_edges']}")
        print(f"    Expansion        : {step7_report['total_expansion_edges']}")
        print(f"  Cross-slide edges  : {len(cross_slide)}")
        print(f"  SRS pool size      : {len(srs_pool)}")
        print(f"\n  Relation distribution:")
        for row in rel_dist_s8:
            print(f"    {row['relation']:<25} {row['count']:>3}  {'█' * row['count']}")
        print(f"\n  Constraint checks:")
        print(f"    1. Slide coverage      : {'PASS' if not uncovered else 'WARN'}")
        print(f"    2. Concept completeness: {'PASS' if not missing else 'WARN'}")
        print(f"    3. Evidence provenance : {'PASS' if n_without_ev == 0 else 'WARN'}")
        print(f"    4. Cross-slide edges   : PASS ({len(cross_slide)} promoted)")

        # ── Save ──────────────────────────────────────────────────────────────────
        self._save_json(lm_edkg, "step8_lm_edkg")

        summary = {
            "doc_id": self.doc_id,
            "pipeline": "PACE-KG-v1",
            "total_slides": step7_report["total_slides"],
            "total_concepts": len(nodes_export),
            "total_edges": len(edges_export),
            "extraction_edges": step7_report["total_extraction_edges"],
            "expansion_edges": step7_report["total_expansion_edges"],
            "cross_slide_edges": len(cross_slide),
            "srs_pool_size": len(srs_pool),
            "relation_distribution": {
                row["relation"]: row["count"] for row in rel_dist_s8
            },
            "constraint_checks": {
                "slide_coverage": len(uncovered) == 0,
                "concept_completeness": len(missing) == 0,
                "evidence_provenance": n_without_ev == 0,
                "cross_slide_promoted": True,
            },
        }

        self._save_json(summary, "step8_summary")

        return lm_edkg

    def step9_generate_summaries(
        self, slides: List[SlideContent], lm_edkg: dict
    ) -> List[SlideSummary]:
        """Generate slide-level summaries."""
        print(f"[Step 9] Generating summaries")

        if self.llm_summary is None:
            self.llm_summary = ChatGroq(
                model=settings.llm_summary,
                temperature=0,
                api_key=settings.groq_api_key,
            )

        summaries: List[SlideSummary] = []

        for sc in slides:
            # Get slide's concepts and triples
            slide_nodes = [
                n
                for n in lm_edkg.get("nodes", [])
                if sc.slide_id in n.get("slide_ids", [])
            ]
            slide_edges = [
                e for e in lm_edkg.get("edges", []) if e.get("slide_id") == sc.slide_id
            ]

            heading = sc.headings[0] if sc.headings else sc.slide_id
            key_terms = [n["name"] for n in slide_nodes]

            # If no triples, use placeholder
            if not slide_edges:
                summary = SlideSummary(
                    slide_id=sc.slide_id,
                    page_number=sc.page_number,
                    heading=heading,
                    summary="This slide contains a diagram — refer to the original slides.",
                    key_terms=key_terms,
                    doc_id=sc.doc_id,
                )
                summaries.append(summary)
                continue

            # Build prompt
            relations_str = "\n".join(
                [f"- {e['subject']} {e['relation']} {e['object']}" for e in slide_edges]
            )

            prompt = f"""SLIDE TITLE: {heading}

KEY CONCEPTS: {", ".join(key_terms)}

RELATIONSHIPS:
{relations_str}

SLIDE TEXT:
{sc.clean_text}

Write a 2-3 sentence summary a student can use as revision notes.
Do not use bullet points. Write in plain English prose."""

            try:
                resp = self.llm_summary.invoke([HumanMessage(content=prompt)])
                summary_text = resp.content.strip()
            except Exception as e:
                print(f"  Summary generation failed for {sc.slide_id}: {e}")
                summary_text = "Summary generation failed."

            summary = SlideSummary(
                slide_id=sc.slide_id,
                page_number=sc.page_number,
                heading=heading,
                summary=summary_text,
                key_terms=key_terms,
                doc_id=sc.doc_id,
            )
            summaries.append(summary)

        # Save JSON
        data = [asdict(s) for s in summaries]
        self._save_json(data, "step9_summaries")

        # Generate Word document
        self._generate_word_doc(summaries)

        print(f"[Step 9] Generated {len(summaries)} summaries")
        return summaries

    def _generate_word_doc(self, summaries: List[SlideSummary]):
        """Export summaries to Word document."""
        doc = Document()

        for summary in summaries:
            # Heading
            heading = doc.add_heading(summary.heading, level=1)
            heading.runs[0].font.size = Pt(14)

            # Summary
            para = doc.add_paragraph(summary.summary)
            para.runs[0].font.size = Pt(11)

            # Key terms
            terms_para = doc.add_paragraph(f"Key terms: {', '.join(summary.key_terms)}")
            terms_para.runs[0].font.italic = True
            terms_para.runs[0].font.size = Pt(11)

            # Page break
            doc.add_page_break()

        # Save
        docx_path = self.output_dir / f"{self.doc_id}_summaries.docx"
        doc.save(str(docx_path))
        print(f"  Word document saved: {docx_path}")

    def run_full_pipeline(self):
        """Run all 9 steps sequentially."""
        print(f"\n{'=' * 80}")
        print(f"Starting PACE-KG pipeline for doc_id: {self.doc_id}")
        print(f"{'=' * 80}\n")

        try:
            # Step 1
            slides_md = self.step1_parse_pdf()

            # Step 2
            slides_content = self.step2_preprocess(slides_md)

            # Step 3
            keyphrases = self.step3_extract_keyphrases(slides_content)

            # Step 4
            triples = self.step4_extract_triples(slides_content, keyphrases)

            # Step 5
            concepts, pruned_triples = self.step5_weight_and_prune(
                triples, keyphrases, slides_content
            )

            # Step 6
            expansions = self.step6_expand(
                concepts, pruned_triples, slides_content, keyphrases
            )

            # Step 7
            self.step7_store_neo4j(concepts, pruned_triples, expansions, slides_content)

            # Step 8
            lm_edkg = self.step8_aggregate(slides_content)

            # Step 9
            summaries = self.step9_generate_summaries(slides_content, lm_edkg)

            print(f"\n{'=' * 80}")
            print(f"Pipeline completed successfully!")
            print(f"{'=' * 80}\n")

        except Exception as e:
            print(f"\n{'=' * 80}")
            print(f"Pipeline failed: {e}")
            print(f"{'=' * 80}\n")
            raise
